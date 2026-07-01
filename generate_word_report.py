import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report():
    output_docx = "g:\\RNG Extractors\\Report\\RNG_Extractors_Comprehensive_Report.docx"
    
    document = Document()
    
    # Define styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    def add_heading(text, level):
        heading = document.add_heading(text, level=level)
        heading.style.font.name = 'Arial'

    def add_paragraph(text):
        # We need to handle basic bold tags since python-docx doesn't parse HTML in add_paragraph
        # A simple parser for <b>...</b>
        p = document.add_paragraph()
        parts = text.split('<b>')
        for part in parts:
            if '</b>' in part:
                bold_text, rest = part.split('</b>', 1)
                p.add_run(bold_text).bold = True
                if rest:
                    p.add_run(rest)
            else:
                p.add_run(part)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15

    def add_image(filepath, width_inches=6.0):
        if os.path.exists(filepath):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(filepath, width=Inches(width_inches))
        else:
            p = document.add_paragraph()
            p.add_run(f"[Image not found: {filepath}]").italic = True

    # Title
    add_heading("RNG Extractors - Web Application Comprehensive Report", 0)

    # 1. Introduction & Workflow
    add_heading("1. User Workflow and Extracted Output", 1)
    add_paragraph("The RNG Extractors web application is an advanced academic platform engineered to process raw bit streams from random number generators. It applies 20 specialized software extraction algorithms to enhance the randomness and strictly evaluates them against the NIST SP 800-22 cryptographic test suite. The platform serves as an end-to-end architecture pipeline for debiasing and entropy extraction.")
    add_paragraph("<b>General User Workflow:</b>")
    add_paragraph("<b>Step 1: Data Upload:</b> The user initiates the process by uploading raw bit streams in either a packed binary (.bin) or plain text (.txt) format onto the platform.")
    add_paragraph("<b>Step 2: Algorithm Selection:</b> From a categorized list of 20 algorithms (Debiasing, Hash-Based, Matrix/Algebra, Mixing Methods, etc.), the user selects which extractors to run on the uploaded bits.")
    add_paragraph("<b>Step 3: Pipeline Execution:</b> By triggering the execution, the platform asynchronously processes the algorithms and subsequently conducts the NIST SP 800-22 statistical test battery to evaluate the cryptographic viability of the extracted streams.")
    add_paragraph("<b>Step 4: Analytics Review:</b> A comprehensive analytics dashboard is generated, offering interactive charts to compare computational efficiency, entropy, throughput, bias levels, and overall NIST pass rates across the chosen methods.")
    add_paragraph("<b>Step 5: Exporting Results:</b> The user can export the processed, high-quality bit streams into a ZIP archive for downstream cryptographic applications and generate a professional PDF report summarizing the comparative analysis.")

    # 2. Overview Page
    document.add_page_break()
    add_heading("2. Overview Page", 1)
    add_paragraph("The Overview page serves as the entry point, featuring a modern, vibrant aesthetic. It provides a high-level summary of the tool's capabilities: processing through 20 extraction methods, evaluated against 15 NIST statistical tests, culminating in live comparative reports.")
    add_paragraph("The page details the Complete Architecture Pipeline via an illustrative flowchart. Key buttons include 'Start Analysis', which navigates the user directly to the primary interaction pipeline (highlighted by a red arrow), and 'Read Documentation' for technical guidance.")
    add_image(r"g:\RNG Extractors\Report\Overview.png", 6.0)

    # 3. Documentation Page
    document.add_page_break()
    add_heading("3. Documentation Page", 1)
    add_paragraph("Before beginning the analysis, users are advised to consult the Documentation page, which outlines system requirements, input specifications, and performance limits to ensure optimal usage.")
    add_paragraph("<b>Input Formats:</b> The system accepts packed binary files (.bin), highly recommended as they are 8x smaller, facilitating faster uploads and processing. Plain text files (.txt) are also supported for legacy compatibility but consume more memory.")
    add_paragraph("<b>System Performance Tiers:</b> The documentation explicitly categorizes file sizes into tiers to set user expectations regarding processing time, especially given the computational intensity of certain O(n²) algorithms:")
    add_paragraph("• <b>Tier 1 (Comprehensive Analysis - Optimal):</b> For inputs ≤ 6.1 KB. Execution takes under 30 seconds.")
    add_paragraph("• <b>Tier 2 (Fast Path Only - Restricted):</b> For inputs between 6.1 KB and 10.0 MB. Execution ranges around 1-2 minutes.")
    add_paragraph("• <b>Tier 3 (Extended Processing - Heavy Load):</b> For large files > 10.0 MB. Extended wait times exceeding 5 minutes.")
    add_image(r"g:\RNG Extractors\Report\Doumentation.png", 6.0)

    # 4. Analyze Page Workflow
    document.add_page_break()
    add_heading("4. Analyze Page and Interaction Flow", 1)
    add_paragraph("The Analyze page is the core functional area of the application. It consists of a multi-step pipeline for uploading data, configuring parameters, and interacting with analytical results.")
    
    add_heading("Step 4.1: Data Input and Configuration", 2)
    add_paragraph("The user encounters a clean interface with a drag-and-drop 'Input Source' zone on the left. The prominent red arrow highlights where users must upload their raw bits (.bin or .txt). Once uploaded, the file details (name and size) populate this block. On the right, the user can select individual or all 20 Extraction Algorithms to be applied.")
    add_image(r"g:\RNG Extractors\Report\Analyze 1.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 2.png", 6.0)

    document.add_page_break()
    add_heading("Step 4.2: Execution and Logging", 2)
    add_paragraph("With the file loaded and algorithms selected, the user clicks the 'Execute Pipeline' button (indicated by the red arrow). The application seamlessly transitions into a live 'Execution Log', displaying timestamped console outputs detailing batch processing, pipeline initialization, algorithm execution, NIST battery testing, and report compilation.")
    add_image(r"g:\RNG Extractors\Report\Analyze 3.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 4.png", 6.0)

    document.add_page_break()
    add_heading("Step 4.3: Analysis Complete & Main Dashboard", 2)
    add_paragraph("Upon completion, the Analysis Complete banner provides a unique Session ID and action buttons for export. The dashboard below visualizes the data through interactive charts mapping Computational Efficiency (in milliseconds), Entropy Comparison, NIST SP 800-22 Pass Rate, Throughput (Bits Per Second), and Bias Level. A dynamic banner showcases the 'Optimal Method', and a tabular Performance Rankings list offers precise metric values for sorting and rigorous inspection.")
    add_image(r"g:\RNG Extractors\Report\Analyze 5.png", 5.5)

    document.add_page_break()
    add_heading("Step 4.4: Interactive Chart Exploration", 2)
    add_paragraph("The charts are highly interactive, allowing users to explore the data in depth. When a user clicks on the bars within a chart, the graph gets expanded for a more detailed, focused view. Furthermore, hovering over any specific bar in the expanded graph reveals a tooltip displaying the exact name of the extraction algorithm along with its corresponding performance metrics, which is crucial for deep cryptographic review.")
    add_paragraph("<b>Entropy & Bias:</b> By utilizing this click-to-expand and hover interaction on the Entropy chart, users can clearly see the exact extractor name alongside its Shannon and Min-Entropy values, visualizing how close an extractor gets to the ideal 1.0 limit. Similarly, interacting with the Bias Level chart exposes the specific extractor's exact numerical bias, where lower values are preferred.")
    add_image(r"g:\RNG Extractors\Report\Analyze 6.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 9.png", 6.0)
    
    document.add_page_break()
    add_paragraph("<b>Throughput & NIST Details:</b> Throughput tooltips highlight the Bit Rate (bps) to compare computational efficiency. For the NIST SP 800-22 Pass Rate chart, hovering displays a breakdown of Pass/Fail/Invalid instances. Additionally, interacting with the chart dynamically updates the 'Detailed NIST Results' side panel, providing a granular test-by-test breakdown (e.g., Frequency, Runs, Linear Complexity) for either the Raw baseline or a specifically selected algorithm like Elias Debiasing.")
    add_image(r"g:\RNG Extractors\Report\Analyze 10.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 7.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 8.png", 6.0)

    document.add_page_break()
    add_heading("Step 4.5: Data Export and Archiving", 2)
    add_paragraph("The platform facilitates the extraction of the processed data for external use. Clicking the 'Export Bitstreams' button triggers an 'Archiving...' state. The user is prompted via a native file dialog to save a ZIP file containing the independent text files for every selected extraction algorithm alongside the raw baseline.")
    add_image(r"g:\RNG Extractors\Report\Analyze 11.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 13.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 14.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 15.png", 6.0)

    document.add_page_break()
    add_heading("Step 4.6: PDF Report Generation", 2)
    add_paragraph("For comprehensive record-keeping, the user clicks the 'Generate PDF' button. After a brief 'Compiling...' phase, a multi-page PDF document is produced, formally summarizing the comparative analysis, visual charts, and ranked tabular metrics.")
    add_image(r"g:\RNG Extractors\Report\Analyze 12.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 16.png", 6.0)
    add_image(r"g:\RNG Extractors\Report\Analyze 17.png", 6.0)

    # 5. About Page
    document.add_page_break()
    add_heading("5. About Page", 1)
    add_paragraph("The About page provides essential academic research background, establishing the credibility and origins of the project. The Project Abstract outlines the core objective: evaluating and enhancing RNGs using software algorithms to secure modern cryptography against predictive attacks.")
    add_paragraph("It proudly highlights the affiliation with the NED University of Engineering and Technology (Department of Physics & Centre for Quantum Technologies). Finally, it acknowledges the dedicated Research Team driving the initiative, including supervisor Dr. Roohi Zafar, co-supervisors and co-PIs, CQT Dr. Tahir Malik and Dr. Kamran, and core developer Syed Muhammad Saad Hussain Zaidi.")
    add_image(r"g:\RNG Extractors\Report\About.png", 6.0)

    document.save(output_docx)
    print("Word Report generated successfully.")

if __name__ == '__main__':
    create_word_report()
